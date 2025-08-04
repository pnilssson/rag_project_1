import fitz
import pytesseract
from PIL import Image
from lxml import etree
import logging
from typing import Optional
from pathlib import Path
import os

# Try to import docx support
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not installed. DOCX files will not be supported.")

from scripts.config import config
from scripts.utils import get_file_extension, is_supported_file, setup_logging

logger = setup_logging()

def extract_text_from_pdf(path: str) -> str:
    """Extract text from PDF with better error handling"""
    try:
        doc = fitz.open(path)
        text_parts = []
        
        for page_num, page in enumerate(doc):
            try:
                # Try to get text first
                text = page.get_text()
                if text.strip():
                    text_parts.append(text)
                else:
                    # If no text, try OCR
                    logger.info(f"Page {page_num + 1} has no text, attempting OCR...")
                    pix = page.get_pixmap()
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    ocr_text = pytesseract.image_to_string(img)
                    if ocr_text.strip():
                        text_parts.append(ocr_text)
            except Exception as e:
                logger.error(f"Error processing page {page_num + 1}: {e}")
                continue
        
        doc.close()
        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"Error extracting text from PDF {path}: {e}")
        raise

def extract_text_from_image(path: str) -> str:
    """Extract text from image using OCR"""
    try:
        image = Image.open(path)
        # Convert to RGB if necessary
        if image.mode != 'RGB':
            image = image.convert('RGB')
        
        # Try different OCR configurations
        text = pytesseract.image_to_string(image, lang='eng+swe')
        if not text.strip():
            # Try without language specification
            text = pytesseract.image_to_string(image)
        
        return text
    except Exception as e:
        logger.error(f"Error extracting text from image {path}: {e}")
        raise

def extract_text_from_xml(path: str) -> str:
    """Extract text from XML with better formatting"""
    try:
        tree = etree.parse(path)
        # Try to extract meaningful text content
        text_elements = tree.xpath('//text()')
        text_content = ' '.join([elem.strip() for elem in text_elements if elem.strip()])
        
        if not text_content:
            # Fallback to full XML as string
            text_content = etree.tostring(tree, pretty_print=True, encoding='unicode')
        
        return text_content
    except Exception as e:
        logger.error(f"Error extracting text from XML {path}: {e}")
        raise

def extract_text_from_docx(path: str) -> str:
    """Extract text from DOCX files"""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx is required for DOCX support")
    
    try:
        doc = Document(path)
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"Error extracting text from DOCX {path}: {e}")
        raise

def extract_text_from_txt(path: str) -> str:
    """Extract text from plain text files with encoding detection"""
    try:
        # Try UTF-8 first
        with open(path, "r", encoding="utf-8") as f:
            return f.read()
    except UnicodeDecodeError:
        try:
            # Try with different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(path, "r", encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
        except Exception as e:
            logger.error(f"Error reading text file {path}: {e}")
            raise
    except Exception as e:
        logger.error(f"Error extracting text from TXT {path}: {e}")
        raise

def extract_text(path: str) -> str:
    """
    Main extraction function with improved file type detection and error handling.
    
    Args:
        path: Path to the file to extract text from
        
    Returns:
        Extracted text content
        
    Raises:
        ValueError: If file format is not supported
        Exception: For other extraction errors
    """
    if not os.path.exists(path):
        raise FileNotFoundError(f"File not found: {path}")
    
    if not is_supported_file(path, config.supported_extensions):
        raise ValueError(f"Unsupported file format: {get_file_extension(path)}")
    
    logger.info(f"Extracting text from: {path}")
    
    ext = get_file_extension(path)
    
    try:
        if ext == ".pdf":
            return extract_text_from_pdf(path)
        elif ext in [".png", ".jpg", ".jpeg", ".tiff", ".bmp"]:
            return extract_text_from_image(path)
        elif ext == ".xml":
            return extract_text_from_xml(path)
        elif ext == ".docx":
            return extract_text_from_docx(path)
        elif ext == ".txt":
            return extract_text_from_txt(path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")
    except Exception as e:
        logger.error(f"Failed to extract text from {path}: {e}")
        raise
